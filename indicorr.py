from dataclasses import dataclass
import pandas as pd
import numpy as np
import os
import math
import scipy
import matplotlib.pyplot as plt
import re
import seaborn as sns
import statsmodels.api as sm
import scipy.stats as stats
from sklearn.linear_model import LogisticRegression
from sklearn.model_selection import train_test_split
from sklearn.metrics import confusion_matrix
from sklearn.metrics import classification_report
from sklearn.metrics import accuracy_score
from sklearn import metrics
from sklearn import preprocessing
from sklearn.model_selection import cross_val_score
from docx import Document
from docx2pdf import convert
from docx.shared import Inches
from scipy.stats import shapiro
from scipy.stats import kstest
from scipy.stats import shapiro
from scipy.stats import kstest
from scipy.stats import normaltest
from scipy.stats import anderson
from scipy.stats import ttest_ind
import scipy.stats as stats
from scipy.stats import spearmanr
from scipy.stats import pearsonr
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import statsmodels.api as sm
  
  
  
  
  
  
    # CORRELATION PER PARTICIPANT 
class indicorr():
    
    def __init__(self, dataframe, path):
        self.dataframe = dataframe
        self.path = path
    

    def indi(self):
        dataframe=self.dataframe
        path=self.path

        doc = Document()


        new_dict = {'ID':[], 'SC_BC_E2_Pearson_r':[], 'SC_BC_E2_Pearson_p_value':[], 'SC_BC_P4_Pearson_r':[], 'SC_BC_P4_Pearson_p_value':[], 
        'SF_BC_E2_Pearson_r':[], 'SF_BC_E2_Pearson_p_value':[], 'SF_BC_P4_Pearson_r':[], 'SF_BC_P4_Pearson_p_value':[],
        'SC_SF_E2_Pearson_r':[], 'SC_SF_E2_Pearson_p_value':[], 'SC_SF_P4_Pearson_r':[], 'SC_SF_P4_Pearson_p_value':[], 'SC:BC_E2_Ratio':[], 'SC:BC_P4_Ratio':[], 
        'SC:BC_E2_Ratio_group':[],'SC:BC_P4_Ratio_group':[], 'SC:BC_Ratio_group':[], 'Subgroup':[]}


        #lineplot
        for p in data['ID'].unique():
            data=dataframe.copy()
            data.fillna(0, inplace=True)
            

            #IDlen = range(list(data['ID'].unique()))
            
            data=data.loc[(data['ID']==p)]

            #print(p)
            #print(data)

            if len(data)==0:
                continue
            

            #pearson correlation


            #data.fillna(0, inplace=True)
            #pearson correlation coefficients

            rscbce, pscbce = scipy.stats.pearsonr(data['SC_Estradiol (pg/mL)'], data['BC_Estradiol (pg/mL)'])
            rsfbce, psfbce = scipy.stats.pearsonr(data['SF_Estradiol (pg/mL)'], data['BC_Estradiol (pg/mL)'])
            rsfsce, psfsce = scipy.stats.pearsonr(data['SC_Estradiol (pg/mL)'], data['SF_Estradiol (pg/mL)'])
            rscbcp, pscbcp = scipy.stats.pearsonr(data['SC_Progesterone (pg/mL)'], data['BC_Progesterone (pg/mL)'])
            rsfbcp, psfbcp = scipy.stats.pearsonr(data['SF_Progesterone (pg/mL)'], data['BC_Progesterone (pg/mL)'])
            rsfscp, psfscp = scipy.stats.pearsonr(data['SC_Progesterone (pg/mL)'], data['SF_Progesterone (pg/mL)'])

            data.loc[((data.Round==3.0)) & (data['SC_Estradiol (pg/mL)']==0), 'SC_Estradiol (pg/mL)']=np.nan
            data.loc[((data.Round==3.0)) & (data['SC_Progesterone (pg/mL)']==0), 'SC_Progesterone (pg/mL)']=np.nan
            data.loc[((data.Round==3.0)) & (data['BC_Progesterone (pg/mL)']==0), 'BC_Progesterone (pg/mL)']=np.nan
            data.loc[((data.Round==3.0)) & (data['BC_Estradiol (pg/mL)']==0), 'BC_Estradiol (pg/mL)']=np.nan
            data.loc[((data.Round==3.0)) & (data['SF_Progesterone (pg/mL)']==0), 'SF_Progesterone (pg/mL)']=np.nan
            data.loc[( (data.Round==3.0)) & (data['SF_Estradiol (pg/mL)']==0), 'SF_Estradiol (pg/mL)']=np.nan
            

            fig, (ax, ax3) = plt.subplots(2, figsize=(10,10))
            #plt.subplot(1,2,1)
            #ax.plot(data.Round, data['BC_Estradiol (pg/mL)'], label='Serum Estradiol')

            scax=ax.plot(data.Round, data['SC_Estradiol (pg/mL)'], color='magenta', label='Saliva Clinic')
            sfax=ax.plot(data.Round, data['SF_Estradiol (pg/mL)'], color='orange', label='Saliva Fasting')
            #scax=ax.plot(data.Round, data['SC_Estradiol (pg/mL)_predictedchange'], color='magenta', alpha=0.5, linestyle='dashed')
            #sfax=ax.plot(data.Round, data['SF_Estradiol (pg/mL)_predictedchange'], color='orange', alpha=0.5, linestyle='dashed')
            ax.set_ylabel('Saliva Estradiol (pg/mL)')
            ax.set_title('#{}, Subgroup {}, SBR {}'.format(p, data['Subgroup'].values[0], (data['SC:BC_Ratio_Bucket'].mean())))
            
            ax2=ax.twinx()

            bcax=ax2.plot(data.Round, data['BC_Estradiol (pg/mL)'], color='dodgerblue')
            #bcax=ax2.plot(data.Round, data['BC_Estradiol (pg/mL)_predictedchange'], color='dodgerblue', alpha=0.5, linestyle='dashed')
            ax2.set_xticks([1.0, 2.0, 3.0])
            ax2.set_ylabel('Serum Estradiol (pg/mL)', color='dodgerblue')
            fig.legend(fontsize=12, bbox_to_anchor=(0.12, 0.85), loc='upper left', frameon=False)
            

            #plt.subplot(1, 2, 2)

            scax=ax3.plot(data.Round, data['SC_Progesterone (pg/mL)'], color='magenta', label='Saliva Clinic')

            sfax=ax3.plot(data.Round, data['SF_Progesterone (pg/mL)'], color='orange', label='Saliva Fasting')
            #scax=ax3.plot(data.Round, data['SC_Progesterone (pg/mL)_predictedchange'], color='magenta', alpha=0.5, linestyle='dashed')
            #sfax=ax3.plot(data.Round, data['SF_Progesterone (pg/mL)_predictedchange'], color='orange', alpha=0.5, linestyle='dashed')
            ax3.set_ylabel('Saliva Progesterone (pg/mL)')
            #ax3.set_title('#{}'.format(p))
            
            ax4=ax3.twinx()

            bcax=ax4.plot(data.Round, data['BC_Progesterone (pg/mL)'], color='dodgerblue', label='Serum')
            #bcax=ax4.plot(data.Round, data['BC_Progesterone (pg/mL)_predictedchange'], color='dodgerblue', alpha=0.5, linestyle='dashed')
            ax4.set_xticks([1.0, 2.0, 3.0])
            ax4.set_ylabel('Serum Progesterone (pg/mL)', color='dodgerblue')
            #fig.legend(fontsize=10, bbox_to_anchor=(0.12, 0.85), loc='upper left', frameon=False)
            
            image = os.path.join(path+"Figures/"+"linegraphperparticipant_#"+str(p)+".png")
            plt.savefig(image)
            plt.show()

            new_dict['ID'].append(p)
            new_dict['SC_BC_E2_Pearson_r'].append(rscbce)
            new_dict['SC_BC_E2_Pearson_p_value'].append(pscbce)
            new_dict['SC_BC_P4_Pearson_r'].append(rscbcp)
            new_dict['SC_BC_P4_Pearson_p_value'].append(pscbcp)
            new_dict['SF_BC_E2_Pearson_r'].append(rsfbce)
            new_dict['SF_BC_E2_Pearson_p_value'].append(psfbce)
            new_dict['SF_BC_P4_Pearson_r'].append(rsfbcp)
            new_dict['SF_BC_P4_Pearson_p_value'].append(psfbcp)
            new_dict['SC_SF_E2_Pearson_r'].append(rsfsce)
            new_dict['SC_SF_E2_Pearson_p_value'].append(psfsce)
            new_dict['SC_SF_P4_Pearson_r'].append(rsfscp)
            new_dict['SC_SF_P4_Pearson_p_value'].append(psfscp)
            new_dict['SC:BC_E2_Ratio'].append(data['SC_Estradiol (pg/mL)'].mean()/data['BC_Estradiol (pg/mL)'].mean())
            new_dict['SC:BC_P4_Ratio'].append(data['SC_Progesterone (pg/mL)'].mean()/data['BC_Progesterone (pg/mL)'].mean())
            new_dict['SC:BC_E2_Ratio_group'].append(data['SC:BC_E2_Ratio_Bucket'].mean())
            new_dict['SC:BC_P4_Ratio_group'].append(data['SC:BC_P4_Ratio_Bucket'].mean())
            new_dict['SC:BC_Ratio_group'].append(data['SC:BC_Ratio_Bucket'].mean())
            new_dict['Subgroup'].append(data['Subgroup'].values[0])
            
            # bucket ratio into third quantiles
            
            

            
            pa = doc.add_paragraph()

            r = pa.add_run()
            
            doc.add_heading('Participant #{}'.format(p), 1)
            doc.add_heading('Response group: {}'.format(data['Subgroup'].values[0]), 2)
            doc.add_heading('Overall Saliva:Blood Serum Ratio (SBR): {}'.format(data['SC:BC_Ratio_Bucket'].mean()), 2)
            doc.add_heading('Estradiol Saliva: Blood Serum ratio group [S1, S2, S3]: {}'.format(data['SC:BC_E2_Ratio_Bucket'].values), 3)
            doc.add_heading('Progesterone Saliva: Blood Serum ratio group [S1, S2, S3] : {}'.format(data['SC:BC_P4_Ratio_Bucket'].values), 3)

            p=doc.add_paragraph()
            r=p.add_run('1 = low secretor, 2 = medium secretor, 3 = high secretor').italic = True
            pa = doc.add_paragraph('    ')
            r = pa.add_run()
            r.add_picture(image, width=Inches(5))
            pa = doc.add_heading('Estradiol', 2)
            
            pa = doc.add_paragraph()
            r = pa.add_run()
            r.add_text('Saliva clinic to Serum Pearson coefficient (r): {}'.format(rscbce))
            r.add_text(' , and p-value (P): {}'.format(pscbce))
            pa = doc.add_paragraph()
            r = pa.add_run()
            r.add_text('Saliva Home to Serum Pearson coefficient (r): {}'.format(rsfbce))
            r.add_text(' , and p-value (P): {}'.format(psfbce))
            pa = doc.add_paragraph()
            r = pa.add_run()
            r.add_text('Saliva Home to Saliva clinic Pearson coefficient (r): {}'.format(rsfsce))
            r.add_text(' , and p-value (P): {}'.format(psfsce))

            pa = doc.add_paragraph(' ')
            pa = doc.add_heading('Progesterone', 2)
            
            pa = doc.add_paragraph()
            r = pa.add_run()
            r.add_text('Saliva clinic to Serum Pearson coefficient (r): {}'.format(rscbcp))
            r.add_text(' , and p-value (P): {}'.format(pscbcp))
            pa = doc.add_paragraph()
            r = pa.add_run()
            r.add_text('Saliva Home to Serum Pearson coefficient (r): {}'.format(rsfbcp))
            r.add_text(' , and p-value (P): {}'.format(psfbcp))
            pa = doc.add_paragraph()
            
            r = pa.add_run()
            r.add_text('Saliva Home to Saliva clinic Pearson coefficient (r): {}'.format(rsfscp))
            r.add_text(' , and p-value (P): {}'.format(psfscp))

            
        correlation_df = pd.DataFrame(new_dict)
        #correlation_df.to_csv(path+"Stats/Individual_correlation_spreadsheet.csv")
        doc.save(path+"Stats/Individual_correlation_stats.docx")

        convert(path+"Stats/Individual_correlation_stats.docx")

        return correlation_df, doc

    def describe(self, data, corrdata):

        ##descriptive statistics

        #data = self.data

        ##histogram of correlation coefficients

        data = data.filter(['SC_Estradiol (pg/mL)', 'BC_Estradiol (pg/mL)', 'SC_Progesterone (pg/mL)', 'BC_Progesterone (pg/mL)', 'SF_Estradiol (pg/mL)', 'SF_Progesterone (pg/mL)', 
        'SC:BC_E2_Ratio', 'SC:BC_E2_Ratio_Bucket', 'SC:BC_P4_Ratio_Bucket', 'SC:BC_Ratio_Bucket', 'Subgroup#'], axis=1)

        #corr heatmap


        sns.heatmap(data.corr(), annot = True)

        

        


